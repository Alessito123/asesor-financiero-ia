from __future__ import annotations

import json
from datetime import datetime
from pathlib import Path

import pandas as pd
from docx import Document
from fpdf import FPDF
from fpdf.enums import XPos, YPos


ROOT = Path(__file__).resolve().parents[1]
OUTPUTS = ROOT / "outputs"
REPORTS = ROOT / "docs"


def load_metrics() -> pd.DataFrame:
    path = OUTPUTS / "model_comparison.csv"
    if path.exists():
        return pd.read_csv(path)
    return pd.DataFrame(
        columns=[
            "model",
            "accuracy_mean",
            "precision_mean",
            "recall_mean",
            "f1_mean",
            "roc_auc_mean",
            "fit_seconds_mean",
        ]
    )


def export_excel(df: pd.DataFrame) -> Path:
    path = REPORTS / "reporte_modelos.xlsx"
    with pd.ExcelWriter(path, engine="openpyxl") as writer:
        df.to_excel(writer, sheet_name="Comparacion", index=False)
        if (OUTPUTS / "fold_results.csv").exists():
            pd.read_csv(OUTPUTS / "fold_results.csv").to_excel(writer, sheet_name="Folds", index=False)
        if (OUTPUTS / "statistical_tests.csv").exists():
            pd.read_csv(OUTPUTS / "statistical_tests.csv").to_excel(writer, sheet_name="Pruebas", index=False)
    return path


def export_word(df: pd.DataFrame) -> Path:
    path = REPORTS / "reporte_modelos.docx"
    doc = Document()
    doc.add_heading("Reporte de modelos de IA", level=1)
    doc.add_paragraph(f"Generado: {datetime.now():%Y-%m-%d %H:%M}")
    doc.add_paragraph(
        "Este reporte resume la comparacion de modelos neuronales clasicos e hibridos "
        "para la prediccion de riesgo financiero."
    )
    if df.empty:
        doc.add_paragraph("No existen metricas. Ejecutar primero ml.training_pipeline.")
    else:
        table = doc.add_table(rows=1, cols=len(df.columns))
        table.style = "Table Grid"
        for idx, column in enumerate(df.columns):
            table.rows[0].cells[idx].text = str(column)
        for _, row in df.iterrows():
            cells = table.add_row().cells
            for idx, column in enumerate(df.columns):
                value = row[column]
                cells[idx].text = f"{value:.4f}" if isinstance(value, float) else str(value)
    doc.save(path)
    return path


def export_pdf(df: pd.DataFrame) -> Path:
    path = REPORTS / "reporte_modelos.pdf"
    pdf = FPDF()
    pdf.add_page()
    pdf.set_font("helvetica", "B", 16)
    pdf.cell(0, 10, "Reporte de modelos de IA", new_x=XPos.LMARGIN, new_y=YPos.NEXT)
    pdf.set_font("helvetica", "", 10)
    pdf.cell(0, 8, f"Generado: {datetime.now():%Y-%m-%d %H:%M}", new_x=XPos.LMARGIN, new_y=YPos.NEXT)
    pdf.ln(4)

    if df.empty:
        pdf.multi_cell(0, 7, "No existen metricas. Ejecutar primero ml.training_pipeline.")
    else:
        for _, row in df.iterrows():
            pdf.set_font("helvetica", "B", 11)
            pdf.cell(0, 7, str(row.get("model", "Modelo")), new_x=XPos.LMARGIN, new_y=YPos.NEXT)
            pdf.set_font("helvetica", "", 10)
            for key in ["accuracy_mean", "precision_mean", "recall_mean", "f1_mean", "roc_auc_mean", "fit_seconds_mean"]:
                if key in row:
                    pdf.cell(0, 6, f"{key}: {row[key]}", new_x=XPos.LMARGIN, new_y=YPos.NEXT)
            pdf.ln(2)
    pdf.output(str(path))
    return path


def main() -> None:
    REPORTS.mkdir(parents=True, exist_ok=True)
    df = load_metrics()
    paths = {
        "excel": str(export_excel(df)),
        "word": str(export_word(df)),
        "pdf": str(export_pdf(df)),
    }
    print(json.dumps(paths, indent=2))


if __name__ == "__main__":
    main()
