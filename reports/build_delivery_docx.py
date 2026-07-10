from __future__ import annotations

from pathlib import Path
from typing import List

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
SOURCE = ROOT / "docs" / "entrega_final.md"
OUTPUT = ROOT / "docs" / "entrega_final_asesor_financiero_ia.docx"


def configure(doc: Document) -> None:
    section = doc.sections[0]
    section.top_margin = Pt(72)
    section.bottom_margin = Pt(72)
    section.left_margin = Pt(72)
    section.right_margin = Pt(72)

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(10.5)
    normal.paragraph_format.space_after = Pt(6)

    for style_name, size in [("Heading 1", 17), ("Heading 2", 13)]:
        style = doc.styles[style_name]
        style.font.name = "Calibri"
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor(31, 77, 120)
        style.paragraph_format.space_before = Pt(10)
        style.paragraph_format.space_after = Pt(6)


def parse_table(lines: List[str]) -> List[List[str]]:
    rows = []
    for line in lines:
        values = [cell.strip() for cell in line.strip().strip("|").split("|")]
        if values and not all(set(value) <= {"-", " "} for value in values):
            rows.append(values)
    return rows


def add_table(doc: Document, rows: List[List[str]]) -> None:
    if not rows:
        return

    width = max(len(row) for row in rows)
    table = doc.add_table(rows=1, cols=width)
    table.style = "Table Grid"
    for index, value in enumerate(rows[0]):
        table.rows[0].cells[index].text = value
        for run in table.rows[0].cells[index].paragraphs[0].runs:
            run.bold = True

    for row in rows[1:]:
        cells = table.add_row().cells
        for index, value in enumerate(row):
            cells[index].text = value


def build_document() -> None:
    doc = Document()
    configure(doc)
    lines = SOURCE.read_text(encoding="utf-8").splitlines()
    table_buffer: List[str] = []
    in_code = False

    def flush_table() -> None:
        nonlocal table_buffer
        if table_buffer:
            add_table(doc, parse_table(table_buffer))
            table_buffer = []

    for raw_line in lines:
        line = raw_line.rstrip()
        if line.startswith("|"):
            table_buffer.append(line)
            continue

        flush_table()
        if line.startswith("```"):
            in_code = not in_code
            continue

        if in_code:
            paragraph = doc.add_paragraph(line)
            paragraph.style = doc.styles["Normal"]
            for run in paragraph.runs:
                run.font.name = "Consolas"
                run.font.size = Pt(9)
            continue

        if line.startswith("# "):
            paragraph = doc.add_heading(line[2:].strip(), level=1)
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        elif line.startswith("## "):
            doc.add_heading(line[3:].strip(), level=2)
        elif line.startswith("- "):
            doc.add_paragraph(line[2:].strip(), style="List Bullet")
        elif line.strip():
            doc.add_paragraph(line.strip())

    flush_table()
    doc.save(OUTPUT)


if __name__ == "__main__":
    build_document()
    print(OUTPUT)
