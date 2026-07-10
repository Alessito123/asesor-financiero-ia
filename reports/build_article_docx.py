from __future__ import annotations

import json
from datetime import date
from pathlib import Path
from typing import Iterable, List

import pandas as pd
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
DOCS = ROOT / "docs"
OUTPUTS = ROOT / "outputs"
MODELS = ROOT / "models"
OUTPUT = DOCS / "articulo_cientifico_asesor_financiero_ia.docx"

BLUE = RGBColor(46, 116, 181)
DARK_BLUE = RGBColor(31, 77, 120)
MUTED = RGBColor(89, 89, 89)
LIGHT_FILL = "F4F6F9"
HEADER_FILL = "E8EEF5"
BORDER = "B7C4D6"


def fmt_decimal(value: object, digits: int = 3) -> str:
    try:
        return f"{float(value):.{digits}f}"
    except (TypeError, ValueError):
        return "Pendiente"


def fmt_seconds(value: object) -> str:
    try:
        return f"{float(value):.2f} s"
    except (TypeError, ValueError):
        return "Pendiente"


def load_metadata() -> dict:
    path = MODELS / "model_metadata.json"
    if path.exists():
        return json.loads(path.read_text(encoding="utf-8"))
    return {}


def comparison_rows() -> List[List[str]]:
    path = OUTPUTS / "model_comparison.csv"
    if not path.exists():
        return [
            ["MLP", "Pendiente", "Pendiente", "Pendiente", "Pendiente", "Pendiente", "Pendiente"],
            ["LSTM", "Pendiente", "Pendiente", "Pendiente", "Pendiente", "Pendiente", "Pendiente"],
            ["GRU", "Pendiente", "Pendiente", "Pendiente", "Pendiente", "Pendiente", "Pendiente"],
            ["CNN-LSTM", "Pendiente", "Pendiente", "Pendiente", "Pendiente", "Pendiente", "Pendiente"],
            ["LSTM-Attention", "Pendiente", "Pendiente", "Pendiente", "Pendiente", "Pendiente", "Pendiente"],
        ]

    df = pd.read_csv(path)
    rows: List[List[str]] = []
    for _, row in df.iterrows():
        rows.append(
            [
                str(row.get("model", "")),
                fmt_decimal(row.get("accuracy_mean")),
                fmt_decimal(row.get("precision_mean")),
                fmt_decimal(row.get("recall_mean")),
                fmt_decimal(row.get("f1_mean")),
                fmt_decimal(row.get("roc_auc_mean")),
                fmt_seconds(row.get("fit_seconds_mean")),
            ]
        )
    return rows


def tuning_rows() -> List[List[str]]:
    path = OUTPUTS / "hyperparameter_tuning.csv"
    if not path.exists():
        return [["Pendiente", "Pendiente", "Pendiente", "Pendiente", "Pendiente"]]

    df = pd.read_csv(path).sort_values("roc_auc", ascending=False)
    rows: List[List[str]] = []
    for _, row in df.iterrows():
        rows.append(
            [
                str(row.get("model", "")),
                fmt_decimal(row.get("learning_rate"), 4),
                str(int(row.get("batch_size", 0))),
                fmt_decimal(row.get("roc_auc")),
                fmt_seconds(row.get("fit_seconds")),
            ]
        )
    return rows


def statistical_rows() -> List[List[str]]:
    path = OUTPUTS / "statistical_tests.csv"
    if not path.exists():
        return [["Pendiente", "Pendiente", "Pendiente", "Pendiente", "Pendiente"]]

    df = pd.read_csv(path)
    rows: List[List[str]] = []
    for _, row in df.iterrows():
        rows.append(
            [
                str(row.get("model", "")),
                str(row.get("baseline", "")),
                fmt_decimal(row.get("paired_t_pvalue"), 4),
                fmt_decimal(row.get("wilcoxon_pvalue"), 4),
                "Significativo" if float(row.get("paired_t_pvalue", 1)) < 0.05 else "No concluyente",
            ]
        )
    return rows


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in {"top": top, "start": start, "bottom": bottom, "end": end}.items():
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths: List[int]) -> None:
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths)))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "120")
    tbl_ind.set(qn("w:type"), "dxa")

    grid = tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        grid_col = OxmlElement("w:gridCol")
        grid_col.set(qn("w:w"), str(width))
        grid.append(grid_col)

    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            cell.width = Inches(widths[idx] / 1440)
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.tcW
            tc_w.set(qn("w:w"), str(widths[idx]))
            tc_w.set(qn("w:type"), "dxa")
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_margins(cell)


def set_font(run, size=None, color=None, bold=None, italic=None) -> None:
    run.font.name = "Calibri"
    run._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    run._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    if size:
        run.font.size = Pt(size)
    if color:
        run.font.color.rgb = color
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def configure_styles(doc: Document) -> None:
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(8)
    normal.paragraph_format.line_spacing = 1.333
    normal.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    for style_name, size, color, before, after in [
        ("Heading 1", 16, BLUE, 18, 10),
        ("Heading 2", 13, BLUE, 12, 6),
        ("Heading 3", 12, DARK_BLUE, 8, 4),
    ]:
        style = doc.styles[style_name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style.font.size = Pt(size)
        style.font.color.rgb = color
        style.font.bold = True
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)


def add_cover(doc: Document) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(120)
    p.paragraph_format.space_after = Pt(16)
    r = p.add_run("Avance de articulo cientifico")
    set_font(r, size=12, color=MUTED, bold=True)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(8)
    r = p.add_run("Prediccion del Riesgo Financiero Personal mediante Modelos Clasicos e Hibridos de Redes Neuronales")
    set_font(r, size=25, color=DARK_BLUE, bold=True)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(28)
    r = p.add_run("Sistema Asesor Financiero Personal IA con Streamlit y FastAPI")
    set_font(r, size=14, color=MUTED, italic=True)

    table = doc.add_table(rows=4, cols=2)
    set_table_geometry(table, [2300, 4760])
    rows = [
        ("Curso / proyecto", "Sistema Asesor Financiero Personal IA"),
        ("Dataset publico", "UCI Default of Credit Card Clients"),
        ("Modelos", "MLP, LSTM, GRU, CNN-LSTM y LSTM-Attention"),
        ("Fecha", date.today().strftime("%d/%m/%Y")),
    ]
    for row_idx, (label, value) in enumerate(rows):
        table.cell(row_idx, 0).text = label
        table.cell(row_idx, 1).text = value
        set_cell_shading(table.cell(row_idx, 0), LIGHT_FILL)
        for cell in table.rows[row_idx].cells:
            for paragraph in cell.paragraphs:
                paragraph.paragraph_format.space_after = Pt(0)
                for run in paragraph.runs:
                    set_font(run, size=10.5, bold=(cell is table.cell(row_idx, 0)))

    doc.add_page_break()


def add_para(doc: Document, text: str) -> None:
    p = doc.add_paragraph(text)
    p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY


def add_bullets(doc: Document, items: Iterable[str]) -> None:
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.line_spacing = 1.208
        p.add_run(item)


def add_numbered(doc: Document, items: Iterable[str]) -> None:
    for item in items:
        p = doc.add_paragraph(style="List Number")
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.line_spacing = 1.208
        p.add_run(item)


def add_table(doc: Document, headers: List[str], rows: List[List[str]], widths: List[int]) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    set_table_geometry(table, widths)
    for idx, header in enumerate(headers):
        cell = table.rows[0].cells[idx]
        cell.text = header
        set_cell_shading(cell, HEADER_FILL)
        for paragraph in cell.paragraphs:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in paragraph.runs:
                set_font(run, size=9.5, bold=True, color=DARK_BLUE)

    for row in rows:
        cells = table.add_row().cells
        for idx, value in enumerate(row):
            cells[idx].text = value
            for paragraph in cells[idx].paragraphs:
                paragraph.paragraph_format.space_after = Pt(0)
                paragraph.paragraph_format.line_spacing = 1.1
                for run in paragraph.runs:
                    set_font(run, size=9)


def add_caption(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run(text)
    set_font(r, size=9, color=MUTED, italic=True)


def build_document() -> None:
    DOCS.mkdir(parents=True, exist_ok=True)
    doc = Document()
    configure_styles(doc)
    metadata = load_metadata()
    final_metrics = metadata.get("final_metrics", {})
    best_model = metadata.get("best_model", "Pendiente")
    add_cover(doc)

    doc.add_heading("Resumen", level=1)
    add_para(
        doc,
        "El presente avance propone el desarrollo y evaluacion de un modulo de inteligencia artificial "
        "para un sistema asesor financiero personal. El objetivo es predecir el riesgo de incumplimiento "
        "financiero de un usuario a partir de variables de credito, historial de pagos y comportamiento "
        "mensual. Para ello se emplea el dataset publico Default of Credit Card Clients del UCI Machine "
        "Learning Repository, una base de datos de 30.000 registros y 23 variables predictoras. La propuesta "
        "contempla tres modelos clasicos de redes neuronales, MLP, LSTM y GRU, y dos modelos hibridos, "
        "CNN-LSTM y LSTM con mecanismo de atencion. El sistema se implementa en Python con Streamlit para "
        "el dashboard y FastAPI para el consumo del mejor modelo entrenado."
    )
    if final_metrics:
        add_para(
            doc,
            "El avance incluye la arquitectura del proyecto, el pipeline reproducible de EDA, entrenamiento, "
            "validacion cruzada de 5 folds, tuning de hiperparametros, pruebas estadisticas robustas y reportes "
            "en PDF, Word y Excel. En la corrida registrada, el mejor modelo fue "
            f"{best_model}, con AUC-ROC final de {fmt_decimal(final_metrics.get('roc_auc'))}, "
            f"accuracy de {fmt_decimal(final_metrics.get('accuracy'))}, precision de "
            f"{fmt_decimal(final_metrics.get('precision'))}, recall de {fmt_decimal(final_metrics.get('recall'))} "
            f"y F1-score de {fmt_decimal(final_metrics.get('f1'))}."
        )
    else:
        add_para(
            doc,
            "El avance incluye la arquitectura del proyecto, el pipeline reproducible de EDA, entrenamiento, "
            "validacion cruzada de 5 folds, tuning de hiperparametros, pruebas estadisticas robustas y reportes "
            "en PDF, Word y Excel. Las metricas finales se completan al ejecutar el pipeline con el dataset "
            "publico en el entorno del equipo."
        )
    add_para(doc, "Palabras clave: asesor financiero, redes neuronales, riesgo crediticio, FastAPI, Streamlit.")

    doc.add_heading("1. Introduccion", level=1)
    add_para(
        doc,
        "La administracion financiera personal requiere herramientas capaces de transformar datos de pago, "
        "credito y consumo en advertencias tempranas. En proyectos academicos, no basta con mostrar graficos "
        "o reglas de recomendacion; es necesario demostrar un modelo de inteligencia artificial entrenado, "
        "validado y consumido por el sistema. Por ello, el sistema Asesor Financiero Personal IA se mejora "
        "con un modulo predictivo que estima el riesgo de incumplimiento financiero."
    )
    add_para(
        doc,
        "El problema se formula como clasificacion binaria: identificar si un cliente presenta riesgo de "
        "incumplimiento en el mes siguiente. Esta salida puede incorporarse al dashboard para producir "
        "alertas, recomendaciones y reportes interpretables despues de validar las credenciales del usuario."
    )

    doc.add_heading("2. Objetivos", level=1)
    doc.add_heading("2.1 Objetivo general", level=2)
    add_para(
        doc,
        "Desarrollar y validar un modulo de inteligencia artificial basado en redes neuronales para predecir "
        "riesgo financiero personal e integrarlo a un sistema web con Streamlit y FastAPI."
    )
    doc.add_heading("2.2 Objetivos especificos", level=2)
    add_numbered(
        doc,
        [
            "Realizar EDA, limpieza de datos y analisis descriptivo del dataset publico seleccionado.",
            "Entrenar tres modelos clasicos y dos modelos hibridos de redes neuronales.",
            "Comparar modelos mediante metricas de clasificacion, curva ROC, matriz de confusion y tiempo de proceso.",
            "Aplicar validacion cruzada de 5 folds configurable y tuning de hiperparametros.",
            "Validar diferencias de desempeno con pruebas estadisticas por fold.",
            "Guardar el mejor modelo en formato .keras y .h5 para consumirlo desde FastAPI sin reentrenar.",
            "Generar reportes en pantalla, PDF, Word y Excel.",
        ],
    )

    doc.add_heading("3. Dataset publico", level=1)
    add_para(
        doc,
        "Se utiliza el dataset Default of Credit Card Clients del UCI Machine Learning Repository. La fuente "
        "reporta 30.000 instancias, 23 variables predictoras, tarea de clasificacion y ausencia de valores "
        "perdidos. El conjunto corresponde a clientes de tarjeta de credito de Taiwan y la variable objetivo "
        "indica incumplimiento de pago en el mes siguiente."
    )
    add_caption(doc, "Tabla 1. Descripcion general del dataset publico.")
    add_table(
        doc,
        ["Elemento", "Descripcion"],
        [
            ["Fuente", "UCI Machine Learning Repository"],
            ["Nombre", "Default of Credit Card Clients"],
            ["Tarea", "Clasificacion binaria"],
            ["Instancias", "30.000"],
            ["Variables", "23 predictoras mas variable objetivo"],
            ["Variable objetivo", "Incumplimiento de pago el mes siguiente"],
            ["Licencia", "Creative Commons Attribution 4.0 International"],
        ],
        [2400, 6960],
    )

    doc.add_heading("4. Metodologia propuesta", level=1)
    add_para(
        doc,
        "La metodologia se organiza como un pipeline reproducible en Python. El archivo "
        "`ml/training_pipeline.py` carga los datos, ejecuta el EDA, entrena los modelos, calcula metricas, "
        "genera figuras y guarda el mejor modelo. La validacion se realiza con StratifiedKFold para preservar "
        "la proporcion de clases en cada fold."
    )
    add_caption(doc, "Tabla 2. Secuencia metodologica exigida para el proyecto.")
    add_table(
        doc,
        ["Fase", "Actividades", "Salida esperada"],
        [
            ["EDA", "Limpieza, estadisticos descriptivos, clases, correlaciones", "CSV de resumen y mapa de calor"],
            ["Entrenamiento", "MLP, LSTM, GRU, CNN-LSTM, LSTM-Attention", "Tabla comparativa por modelo"],
            ["Validacion cruzada", "5 folds configurables", "Resultados por fold"],
            ["Tuning", "Learning rate, batch size y parametros base", "Mejor configuracion"],
            ["Pruebas estadisticas", "t-test pareado y Wilcoxon sobre AUC por fold", "Evidencia de diferencias"],
            ["Reportes", "Pantalla, PDF, Word y Excel", "Archivos en carpeta docs/"],
        ],
        [1700, 4760, 2900],
    )

    doc.add_heading("5. Modelos de inteligencia artificial", level=1)
    add_para(
        doc,
        "Se entrenan cinco arquitecturas de redes neuronales. MLP opera sobre variables tabulares escaladas. "
        "LSTM y GRU tratan la secuencia de variables financieras como una serie ordenada de caracteristicas. "
        "CNN-LSTM combina extraccion local de patrones y memoria temporal. LSTM-Attention agrega un mecanismo "
        "de atencion para ponderar caracteristicas relevantes."
    )
    add_caption(doc, "Tabla 3. Modelos contemplados en el entrenamiento.")
    add_table(
        doc,
        ["Tipo", "Modelo", "Justificacion"],
        [
            ["Clasico", "MLP", "Base neuronal para datos tabulares y relaciones no lineales."],
            ["Clasico", "LSTM", "Captura dependencias en secuencias de pagos y montos mensuales."],
            ["Clasico", "GRU", "Alternativa recurrente con menor complejidad que LSTM."],
            ["Hibrido", "CNN-LSTM", "Extrae patrones locales y luego modela dependencias secuenciales."],
            ["Hibrido", "LSTM-Attention", "Pondera variables o periodos de mayor influencia predictiva."],
        ],
        [1600, 2200, 5560],
    )

    doc.add_heading("6. Evaluacion y validacion", level=1)
    add_para(
        doc,
        "La evaluacion contempla exactitud, precision, recall, F1-score, AUC-ROC y tiempo de entrenamiento. "
        "Tambien se generan matriz de confusion, curva ROC y mapa de calor. El mejor modelo se selecciona "
        "principalmente por AUC-ROC, porque el problema de riesgo financiero requiere estimar probabilidad y "
        "no solo una etiqueta binaria."
    )
    add_caption(doc, "Tabla 4. Plantilla para resultados comparativos generados por el pipeline.")
    add_table(
        doc,
        ["Modelo", "Accuracy", "Precision", "Recall", "F1", "AUC-ROC", "Tiempo medio"],
        comparison_rows(),
        [1700, 1250, 1250, 1250, 1050, 1300, 1560],
    )
    if final_metrics:
        add_para(
            doc,
            "Las figuras requeridas se generaron en la carpeta `outputs/`: mapa de calor de correlaciones, "
            "matriz de confusion y curva ROC. El mejor modelo se guardo como `models/best_model.keras` y "
            "`models/best_model.h5` para ser consumido por FastAPI sin reentrenamiento."
        )
    else:
        add_para(
            doc,
            "Cuando se ejecute `python -m ml.training_pipeline --folds 5 --epochs 20`, esta tabla se completa "
            "a partir de `outputs/model_comparison.csv`. Las figuras requeridas se generan en `outputs/`."
        )

    add_caption(doc, "Tabla 5. Tuning de hiperparametros del mejor modelo.")
    add_table(
        doc,
        ["Modelo", "Learning rate", "Batch size", "AUC-ROC", "Tiempo"],
        tuning_rows(),
        [1900, 1900, 1700, 1700, 2160],
    )

    add_caption(doc, "Tabla 6. Pruebas estadisticas sobre AUC por fold.")
    add_table(
        doc,
        ["Modelo", "Baseline", "p t-test", "p Wilcoxon", "Interpretacion"],
        statistical_rows(),
        [1700, 1700, 1700, 1700, 2560],
    )

    doc.add_heading("7. Arquitectura del sistema", level=1)
    add_para(
        doc,
        "La arquitectura separa el entrenamiento del consumo del modelo. El entrenamiento produce el modelo "
        "guardado y los artefactos de evaluacion. FastAPI carga el modelo y ofrece endpoints para prediccion. "
        "Streamlit valida credenciales, muestra dashboard, permite formular predicciones y descarga el articulo."
    )
    add_caption(doc, "Tabla 7. Componentes del proyecto mejorado.")
    add_table(
        doc,
        ["Componente", "Tecnologia", "Responsabilidad"],
        [
            ["Dashboard", "Streamlit", "Login, visualizacion, prediccion y descarga de reportes."],
            ["Backend", "FastAPI", "Endpoint /predict, /health y /model-info."],
            ["Modelo", "TensorFlow/Keras", "Redes neuronales clasicas e hibridas."],
            ["Persistencia de modelo", ".keras, .h5, joblib", "Evitar reentrenamiento en produccion."],
            ["Despliegue", "Render / Vercel", "API en Render y frontend auxiliar en Vercel."],
            ["Gestion", "GitHub / Jira", "Control de versiones y seguimiento academico."],
        ],
        [1900, 2300, 5160],
    )

    doc.add_heading("8. Implementacion entregada", level=1)
    add_bullets(
        doc,
        [
            "Se preservo el codigo original en `legacy/` y el flujo n8n en `n8n/`.",
            "Se creo `backend/main.py` con FastAPI para consumir el modelo.",
            "Se creo `frontend/streamlit_app.py` con login, formulario de prediccion y vista de modelos.",
            "Se creo `ml/training_pipeline.py` con EDA, 5 modelos, validacion cruzada, tuning y pruebas estadisticas.",
            "Se agregaron `render.yaml`, `Procfile` y `vercel_frontend/` para despliegue.",
            "Se agrego `docs/jira_backlog.md` como base para documentacion en Jira.",
        ],
    )

    doc.add_heading("9. Consideraciones eticas y limitaciones", level=1)
    add_para(
        doc,
        "El modelo estima riesgo a partir de datos historicos y no debe usarse como decision automatica unica. "
        "En un contexto real deben aplicarse criterios de privacidad, explicabilidad, revision humana y control "
        "de sesgos. La version academica usa un dataset publico, por lo que no representa necesariamente el "
        "comportamiento financiero de usuarios peruanos o de todos los segmentos socioeconomicos."
    )

    doc.add_heading("10. Conclusiones preliminares", level=1)
    if final_metrics:
        add_para(
            doc,
            "El proyecto queda reorientado hacia un sistema de inteligencia artificial verificable: posee dataset "
            "publico, modelos definidos, validacion cruzada, tuning, pruebas estadisticas, consumo por API y "
            "dashboard despues del login. La corrida registrada permitio seleccionar formalmente el modelo "
            f"{best_model} como mejor alternativa por AUC-ROC. Para una entrega final mas robusta, se recomienda "
            "repetir el entrenamiento con mayor numero de epocas y comparar la estabilidad de resultados."
        )
    else:
        add_para(
            doc,
            "El proyecto queda reorientado hacia un sistema de inteligencia artificial verificable: posee dataset "
            "publico, modelos definidos, validacion cruzada, tuning, pruebas estadisticas, consumo por API y "
            "dashboard despues del login. El siguiente paso es ejecutar el pipeline de entrenamiento con el dataset "
            "UCI, completar la tabla de resultados, interpretar las figuras generadas y seleccionar formalmente el "
            "mejor modelo para produccion."
        )

    doc.add_heading("Referencias", level=1)
    references = [
        "Yeh, I. (2009). Default of Credit Card Clients [Dataset]. UCI Machine Learning Repository. https://doi.org/10.24432/C55S3H",
        "UCI Machine Learning Repository. Default of Credit Card Clients. https://archive.ics.uci.edu/dataset/350/default+of+credit+card+clients",
        "Yeh, I. C., & Lien, C. H. (2009). The comparisons of data mining techniques for the predictive accuracy of probability of default of credit card clients. Expert Systems with Applications.",
        "Chollet, F. Keras: Deep Learning for humans. https://keras.io/",
        "FastAPI documentation. https://fastapi.tiangolo.com/",
        "Streamlit documentation. https://docs.streamlit.io/",
    ]
    for item in references:
        p = doc.add_paragraph(style="List Number")
        p.paragraph_format.space_after = Pt(4)
        p.add_run(item)

    section = doc.sections[0]
    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = footer.add_run("Sistema Asesor Financiero Personal IA - Avance de articulo cientifico")
    set_font(run, size=8, color=MUTED)

    doc.save(OUTPUT)


if __name__ == "__main__":
    build_document()
    print(OUTPUT)
