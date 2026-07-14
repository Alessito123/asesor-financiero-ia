from __future__ import annotations

import os
import time
import unicodedata
from collections import defaultdict, deque
from datetime import datetime
from io import BytesIO
from typing import Any, Dict, List, Optional

from fastapi import FastAPI, Request
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import JSONResponse, Response
from docx import Document
from docx.shared import Inches
from fpdf import FPDF
from fpdf.enums import XPos, YPos
from openpyxl import Workbook
from openpyxl.chart import BarChart, Reference
from openpyxl.styles import Alignment, Border, Font, PatternFill, Side
from pydantic import BaseModel, Field

from ml.predictor import FinancialRiskPredictor
from ml.schema import DEFAULT_SAMPLE, FEATURE_COLUMNS
from ml.statistical_validation import statistical_validation_summary


class PredictionRequest(BaseModel):
    model_name: str = Field("LSTM", pattern="^(MLP|LSTM|GRU|CNN_LSTM|CNN-LSTM|LSTM_ATTENTION|LSTM-ATTENTION)$")
    LIMIT_BAL: float = Field(..., ge=1, le=10_000_000, description="Monto de credito concedido")
    SEX: int = Field(2, ge=1, le=2)
    EDUCATION: int = Field(2, ge=1, le=4)
    MARRIAGE: int = Field(1, ge=1, le=3)
    AGE: int = Field(..., ge=18, le=100)
    PAY_0: int = Field(0, ge=-2, le=9)
    PAY_2: int = Field(0, ge=-2, le=9)
    PAY_3: int = Field(0, ge=-2, le=9)
    PAY_4: int = Field(0, ge=-2, le=9)
    PAY_5: int = Field(0, ge=-2, le=9)
    PAY_6: int = Field(0, ge=-2, le=9)
    BILL_AMT1: float = Field(0, ge=0, le=10_000_000)
    BILL_AMT2: float = Field(0, ge=0, le=10_000_000)
    BILL_AMT3: float = Field(0, ge=0, le=10_000_000)
    BILL_AMT4: float = Field(0, ge=0, le=10_000_000)
    BILL_AMT5: float = Field(0, ge=0, le=10_000_000)
    BILL_AMT6: float = Field(0, ge=0, le=10_000_000)
    PAY_AMT1: float = Field(0, ge=0, le=10_000_000)
    PAY_AMT2: float = Field(0, ge=0, le=10_000_000)
    PAY_AMT3: float = Field(0, ge=0, le=10_000_000)
    PAY_AMT4: float = Field(0, ge=0, le=10_000_000)
    PAY_AMT5: float = Field(0, ge=0, le=10_000_000)
    PAY_AMT6: float = Field(0, ge=0, le=10_000_000)


class PredictionResponse(BaseModel):
    probability: float
    risk_label: str
    mode: str
    threshold: float
    prediction: int
    explanation: List[str]
    model_name: str
    requested_model: str = "LSTM"
    model_available: bool = False
    model_status: str = ""
    artifact_path: Optional[str] = None


class ChatbotRequest(BaseModel):
    message: str = Field(..., min_length=1, max_length=600)
    language: str = Field("es", pattern="^(es|en)$")
    prediction_context: Optional[Dict[str, Any]] = None


class ChatbotResponse(BaseModel):
    answer: str
    language: str
    topics: List[str]
    sources: List[str] = Field(default_factory=list)
    context_used: bool = False


app = FastAPI(
    title="Asesor Financiero Personal IA API",
    description="API para prediccion de riesgo financiero con modelos de redes neuronales.",
    version="1.0.0",
)

app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

RATE_LIMIT_MAX_REQUESTS = int(os.getenv("RATE_LIMIT_MAX_REQUESTS", "10"))
RATE_LIMIT_WINDOW_SECONDS = int(os.getenv("RATE_LIMIT_WINDOW_SECONDS", "60"))
RATE_LIMITED_PATHS = {
    "/predict",
    "/chatbot",
    "/reports/financial/pdf",
    "/reports/financial/docx",
    "/reports/financial/xlsx",
}
_request_log: Dict[str, deque[float]] = defaultdict(deque)


@app.middleware("http")
async def rate_limit_middleware(request: Request, call_next):
    if request.method == "POST" and request.url.path in RATE_LIMITED_PATHS:
        client_host = request.client.host if request.client else "unknown"
        key = f"{client_host}:{request.url.path}"
        now = time.monotonic()
        bucket = _request_log[key]
        while bucket and now - bucket[0] > RATE_LIMIT_WINDOW_SECONDS:
            bucket.popleft()
        if len(bucket) >= RATE_LIMIT_MAX_REQUESTS:
            return JSONResponse(
                status_code=429,
                content={
                    "detail": "Limite de uso alcanzado. Espera un minuto antes de volver a intentar.",
                    "limit": RATE_LIMIT_MAX_REQUESTS,
                    "window_seconds": RATE_LIMIT_WINDOW_SECONDS,
                },
            )
        bucket.append(now)
    return await call_next(request)


predictor = FinancialRiskPredictor()


MONEY_FIELDS = [
    "LIMIT_BAL",
    "BILL_AMT1",
    "BILL_AMT2",
    "BILL_AMT3",
    "BILL_AMT4",
    "BILL_AMT5",
    "BILL_AMT6",
    "PAY_AMT1",
    "PAY_AMT2",
    "PAY_AMT3",
    "PAY_AMT4",
    "PAY_AMT5",
    "PAY_AMT6",
]

FIELD_LABELS = {
    "LIMIT_BAL": "Credito concedido",
    "SEX": "Sexo",
    "EDUCATION": "Educacion",
    "MARRIAGE": "Estado civil",
    "AGE": "Edad",
    "PAY_0": "Mora mes actual",
    "PAY_2": "Mora mes -2",
    "PAY_3": "Mora mes -3",
    "PAY_4": "Mora mes -4",
    "PAY_5": "Mora mes -5",
    "PAY_6": "Mora mes -6",
    "BILL_AMT1": "Saldo 1",
    "BILL_AMT2": "Saldo 2",
    "BILL_AMT3": "Saldo 3",
    "BILL_AMT4": "Saldo 4",
    "BILL_AMT5": "Saldo 5",
    "BILL_AMT6": "Saldo 6",
    "PAY_AMT1": "Pago 1",
    "PAY_AMT2": "Pago 2",
    "PAY_AMT3": "Pago 3",
    "PAY_AMT4": "Pago 4",
    "PAY_AMT5": "Pago 5",
    "PAY_AMT6": "Pago 6",
}

SEX_LABELS = {1: "Masculino", 2: "Femenino"}
EDUCATION_LABELS = {1: "Posgrado", 2: "Universidad", 3: "Secundaria", 4: "Otros"}
MARRIAGE_LABELS = {1: "Casado", 2: "Soltero", 3: "Otros"}

ACADEMIC_KNOWLEDGE = [
    {
        "id": "dataset_uci_credit_card",
        "keywords": ["dataset", "uci", "datos", "variables", "credit card", "taiwan"],
        "es": (
            "El sistema se basa en el dataset publico UCI Default of Credit Card Clients, "
            "con variables de limite de credito, perfil del cliente, historial de mora, saldos facturados y pagos."
        ),
        "en": (
            "The system is based on the public UCI Default of Credit Card Clients dataset, "
            "with credit limit, client profile, delay history, billed balances and payment variables."
        ),
    },
    {
        "id": "model_catalog",
        "keywords": ["modelo", "model", "lstm", "gru", "mlp", "cnn", "attention", "hibrido", "hybrid"],
        "es": (
            "La investigacion compara tres modelos neuronales clasicos (MLP, LSTM, GRU) y dos hibridos "
            "(CNN-LSTM y LSTM-Attention). El modelo guardado se consume desde FastAPI para evitar reentrenar."
        ),
        "en": (
            "The research compares three classic neural models (MLP, LSTM, GRU) and two hybrid models "
            "(CNN-LSTM and LSTM-Attention). The saved model is consumed from FastAPI to avoid retraining."
        ),
    },
    {
        "id": "statistical_validation",
        "keywords": ["estadistica", "statistical", "wilcoxon", "t-test", "cross", "fold", "validacion"],
        "es": (
            "La validacion incluye cross validation configurable de 5 folds, comparacion de metricas y pruebas "
            "estadisticas pareadas como t-test y Wilcoxon para contrastar diferencias entre modelos."
        ),
        "en": (
            "Validation includes configurable 5-fold cross validation, metric comparison and paired statistical "
            "tests such as t-test and Wilcoxon to contrast model differences."
        ),
    },
    {
        "id": "reports_module",
        "keywords": ["reporte", "report", "pdf", "word", "excel", "descargar", "download"],
        "es": (
            "El modulo de reportes genera PDF, Word y Excel del analisis financiero del programa, incluyendo "
            "resultado, indicadores, recomendaciones y datos ingresados."
        ),
        "en": (
            "The reports module generates PDF, Word and Excel files for the program's financial analysis, "
            "including result, indicators, recommendations and entered data."
        ),
    },
    {
        "id": "operational_xai",
        "keywords": ["explica", "explain", "xai", "shap", "porque", "why", "riesgo", "risk"],
        "es": (
            "La explicabilidad operativa resume factores como mora reciente, cobertura de pagos, uso del credito, "
            "limite aprobado y edad. Es una aproximacion interpretable para acompanar la prediccion neuronal."
        ),
        "en": (
            "Operational explainability summarizes factors such as recent delay, payment coverage, credit usage, "
            "approved limit and age. It is an interpretable approximation that accompanies the neural prediction."
        ),
    },
    {
        "id": "app_purpose",
        "keywords": ["para que", "que hace", "programa", "app", "aplicacion", "sistema", "sirve", "purpose"],
        "es": (
            "La aplicacion sirve para estimar riesgo de incumplimiento financiero, simular escenarios, explicar "
            "factores de riesgo y generar reportes del analisis para un proyecto academico de IA financiera."
        ),
        "en": (
            "The app estimates financial default risk, simulates scenarios, explains risk drivers and generates "
            "analysis reports for an academic financial AI project."
        ),
    },
    {
        "id": "model_recommendation",
        "keywords": ["que modelo", "cual modelo", "recomienda", "debo usar", "seleccionar", "usar modelo", "best model"],
        "es": (
            "Para uso principal se recomienda el modelo guardado en produccion. Para comparacion academica se pueden "
            "seleccionar los cinco modelos y revisar AUC, F1, tiempo de proceso y pruebas estadisticas."
        ),
        "en": (
            "For the main workflow, use the saved production model. For academic comparison, select the five models "
            "and review AUC, F1, processing time and statistical tests."
        ),
    },
    {
        "id": "user_workflow",
        "keywords": ["como uso", "como funciona", "pasos", "usuario", "llenar", "formulario", "workflow"],
        "es": (
            "El flujo recomendado es ingresar perfil del cliente, historial de mora, saldos y pagos; elegir el modelo; "
            "calcular riesgo; revisar factores explicativos; comparar escenarios y descargar reportes."
        ),
        "en": (
            "The recommended workflow is to enter client profile, delay history, balances and payments; choose a model; "
            "calculate risk; review explainability factors; compare scenarios and download reports."
        ),
    },
    {
        "id": "limitations",
        "keywords": ["confiable", "limitacion", "decision", "aprobar", "rechazar", "banco", "limitation"],
        "es": (
            "El resultado es apoyo analitico y academico. No debe usarse como unica decision para aprobar o rechazar "
            "creditos; se debe complementar con politicas internas, datos actualizados y revision humana."
        ),
        "en": (
            "The output is analytical and academic support. It should not be the only decision rule for approving or "
            "rejecting credit; it should be complemented with policies, updated data and human review."
        ),
    },
]


def money(value: float) -> str:
    return f"S/ {float(value):,.0f}".replace(",", " ")


def percent(value: float) -> str:
    return f"{float(value) * 100:.1f}%"


def display_value(key: str, value: float) -> str:
    if key in MONEY_FIELDS:
        return money(value)
    if key == "SEX":
        return SEX_LABELS.get(int(value), str(value))
    if key == "EDUCATION":
        return EDUCATION_LABELS.get(int(value), str(value))
    if key == "MARRIAGE":
        return MARRIAGE_LABELS.get(int(value), str(value))
    return str(int(value)) if float(value).is_integer() else str(value)


def financial_indicators(payload: Dict) -> Dict[str, float]:
    bills = [max(float(payload.get(f"BILL_AMT{i}", 0)), 0.0) for i in range(1, 7)]
    payments = [max(float(payload.get(f"PAY_AMT{i}", 0)), 0.0) for i in range(1, 7)]
    delays = [float(payload.get(key, 0)) for key in ["PAY_0", "PAY_2", "PAY_3", "PAY_4", "PAY_5", "PAY_6"]]
    total_bills = sum(bills)
    total_payments = sum(payments)
    limit_bal = max(float(payload.get("LIMIT_BAL", 1)), 1.0)
    return {
        "total_bills": total_bills,
        "total_payments": total_payments,
        "credit_usage": min(total_bills / (limit_bal * 6.0), 2.0),
        "payment_coverage": total_payments / max(total_bills, 1.0),
        "max_delay": max(delays),
        "average_bill": total_bills / 6.0,
        "average_payment": total_payments / 6.0,
    }


def clamp(value: float, lower: float, upper: float) -> float:
    return max(lower, min(upper, value))


def operational_risk_drivers(payload: Dict) -> List[Dict[str, Any]]:
    indicators = financial_indicators(payload)
    limit_bal = float(payload.get("LIMIT_BAL", 0))
    age = int(float(payload.get("AGE", 0)))
    drivers: List[Dict[str, Any]] = []

    if indicators["max_delay"] > 0:
        drivers.append(
            {
                "factor": "Mora reciente",
                "direction": "increase",
                "impact": round(clamp(0.05 + indicators["max_delay"] * 0.055, 0.05, 0.32), 3),
                "detail": f"{int(indicators['max_delay'])} meses de atraso maximo detectado.",
            }
        )
    if indicators["payment_coverage"] < 0.08:
        drivers.append(
            {
                "factor": "Cobertura de pagos",
                "direction": "increase",
                "impact": round(clamp(0.14 - indicators["payment_coverage"], 0.06, 0.18), 3),
                "detail": f"Cobertura acumulada de {percent(indicators['payment_coverage'])}.",
            }
        )
    elif indicators["payment_coverage"] > 0.18:
        drivers.append(
            {
                "factor": "Cobertura de pagos",
                "direction": "reduce",
                "impact": round(clamp(indicators["payment_coverage"] * 0.34, 0.06, 0.18), 3),
                "detail": f"Cobertura acumulada de {percent(indicators['payment_coverage'])}.",
            }
        )
    if indicators["credit_usage"] > 0.65:
        drivers.append(
            {
                "factor": "Uso de credito",
                "direction": "increase",
                "impact": round(clamp((indicators["credit_usage"] - 0.55) * 0.22, 0.05, 0.20), 3),
                "detail": f"Uso promedio de {percent(indicators['credit_usage'])}.",
            }
        )
    elif indicators["credit_usage"] < 0.25:
        drivers.append(
            {
                "factor": "Uso de credito",
                "direction": "reduce",
                "impact": round(clamp((0.3 - indicators["credit_usage"]) * 0.28, 0.04, 0.12), 3),
                "detail": f"Uso promedio de {percent(indicators['credit_usage'])}.",
            }
        )
    if limit_bal >= 250000:
        drivers.append(
            {
                "factor": "Capacidad aprobada",
                "direction": "reduce",
                "impact": 0.07,
                "detail": f"Limite de credito de {money(limit_bal)}.",
            }
        )
    elif 0 < limit_bal < 80000:
        drivers.append(
            {
                "factor": "Capacidad aprobada",
                "direction": "increase",
                "impact": 0.06,
                "detail": f"Limite de credito de {money(limit_bal)}.",
            }
        )
    if 0 < age < 25:
        drivers.append(
            {"factor": "Madurez financiera", "direction": "increase", "impact": 0.04, "detail": f"Edad: {age} anos."}
        )
    elif age >= 45:
        drivers.append(
            {"factor": "Madurez financiera", "direction": "reduce", "impact": 0.04, "detail": f"Edad: {age} anos."}
        )

    return sorted(drivers, key=lambda item: item["impact"], reverse=True)[:5]


def recommendations(result: Dict, indicators: Dict) -> List[str]:
    probability = float(result["probability"])
    notes = []
    if probability >= 0.70:
        notes.append("Revisar la capacidad de pago antes de aprobar nuevas obligaciones.")
        notes.append("Priorizar reduccion de mora y pagos minimos por encima de nuevos consumos.")
    elif probability >= 0.40:
        notes.append("Mantener seguimiento mensual del saldo y evitar crecimiento de deuda.")
        notes.append("Aumentar pagos para mejorar la cobertura frente al saldo facturado.")
    else:
        notes.append("Mantener el comportamiento actual y conservar pagos puntuales.")
        notes.append("Usar el credito de forma moderada para sostener un perfil saludable.")
    if indicators["credit_usage"] > 0.75:
        notes.append("La utilizacion promedio del credito es elevada; conviene reducir saldos.")
    if indicators["payment_coverage"] < 0.10:
        notes.append("La cobertura de pagos es baja frente al saldo total; revisar presupuesto mensual.")
    if indicators["max_delay"] > 0:
        notes.append("Existe mora historica; corregir atrasos mejora la clasificacion de riesgo.")
    return notes


def report_context(request: PredictionRequest) -> Dict:
    payload = request.model_dump()
    result = predictor.predict(payload, request.model_name)
    indicators = financial_indicators(payload)
    return {
        "payload": payload,
        "result": result,
        "indicators": indicators,
        "drivers": operational_risk_drivers(payload),
        "recommendations": recommendations(result, indicators),
        "generated_at": datetime.now().strftime("%Y-%m-%d %H:%M"),
    }


def build_pdf_report(context: Dict) -> bytes:
    payload = context["payload"]
    result = context["result"]
    indicators = context["indicators"]
    drivers = context["drivers"]

    pdf = FPDF()
    pdf.set_auto_page_break(auto=True, margin=15)
    pdf.add_page()
    pdf.set_fill_color(10, 13, 20)
    pdf.rect(0, 0, 210, 30, "F")
    pdf.set_text_color(255, 255, 255)
    pdf.set_font("helvetica", "B", 16)
    pdf.set_xy(12, 10)
    pdf.cell(0, 9, "Asesor Financiero IA", new_x=XPos.LMARGIN, new_y=YPos.NEXT)
    pdf.set_text_color(16, 24, 40)
    pdf.ln(12)
    pdf.set_font("helvetica", "B", 15)
    pdf.cell(0, 9, "Reporte financiero del programa", new_x=XPos.LMARGIN, new_y=YPos.NEXT)
    pdf.set_font("helvetica", "", 10)
    pdf.cell(0, 7, f"Generado: {context['generated_at']}", new_x=XPos.LMARGIN, new_y=YPos.NEXT)
    pdf.cell(0, 7, f"Modelo: {result['model_name']} | modo: {result['mode']}", new_x=XPos.LMARGIN, new_y=YPos.NEXT)
    pdf.ln(3)

    pdf.set_font("helvetica", "B", 13)
    pdf.cell(0, 8, "Resultado del analisis", new_x=XPos.LMARGIN, new_y=YPos.NEXT)
    pdf.set_font("helvetica", "", 10)
    pdf.cell(0, 7, f"Probabilidad de incumplimiento: {percent(result['probability'])}", new_x=XPos.LMARGIN, new_y=YPos.NEXT)
    pdf.cell(0, 7, f"Nivel de riesgo: {result['risk_label'].upper()}", new_x=XPos.LMARGIN, new_y=YPos.NEXT)
    pdf.cell(0, 7, f"Decision del umbral: {'incumplimiento probable' if result['prediction'] else 'sin incumplimiento probable'}", new_x=XPos.LMARGIN, new_y=YPos.NEXT)
    pdf.ln(2)

    pdf.set_font("helvetica", "B", 13)
    pdf.cell(0, 8, "Indicadores financieros", new_x=XPos.LMARGIN, new_y=YPos.NEXT)
    pdf.set_font("helvetica", "", 10)
    rows = [
        ("Saldos totales", money(indicators["total_bills"])),
        ("Pagos totales", money(indicators["total_payments"])),
        ("Uso promedio de credito", percent(indicators["credit_usage"])),
        ("Cobertura pago / saldo", percent(indicators["payment_coverage"])),
        ("Mora maxima", f"{int(indicators['max_delay'])} meses" if indicators["max_delay"] > 0 else "Sin mora"),
    ]
    for label, value in rows:
        pdf.cell(70, 7, label)
        pdf.cell(0, 7, value, new_x=XPos.LMARGIN, new_y=YPos.NEXT)
    pdf.ln(2)

    pdf.set_font("helvetica", "B", 13)
    pdf.cell(0, 8, "Factores explicativos", new_x=XPos.LMARGIN, new_y=YPos.NEXT)
    pdf.set_font("helvetica", "", 10)
    if drivers:
        for driver in drivers:
            sign = "+" if driver["direction"] == "increase" else "-"
            pdf.multi_cell(
                0,
                6,
                f"- {driver['factor']}: {sign}{percent(driver['impact'])}. {driver['detail']}",
                new_x=XPos.LMARGIN,
                new_y=YPos.NEXT,
            )
    else:
        pdf.multi_cell(0, 6, "El perfil no muestra factores extremos de riesgo.", new_x=XPos.LMARGIN, new_y=YPos.NEXT)
    pdf.ln(2)

    pdf.set_font("helvetica", "B", 13)
    pdf.cell(0, 8, "Recomendaciones", new_x=XPos.LMARGIN, new_y=YPos.NEXT)
    pdf.set_font("helvetica", "", 10)
    for item in context["recommendations"]:
        pdf.set_x(pdf.l_margin)
        pdf.multi_cell(0, 6, f"- {item}", new_x=XPos.LMARGIN, new_y=YPos.NEXT)
    pdf.ln(2)

    pdf.set_font("helvetica", "B", 13)
    pdf.cell(0, 8, "Datos ingresados", new_x=XPos.LMARGIN, new_y=YPos.NEXT)
    pdf.set_font("helvetica", "", 9)
    for key in FEATURE_COLUMNS:
        pdf.cell(62, 6, FIELD_LABELS.get(key, key))
        pdf.cell(0, 6, display_value(key, payload[key]), new_x=XPos.LMARGIN, new_y=YPos.NEXT)
    return bytes(pdf.output(dest="S"))


def build_docx_report(context: Dict) -> bytes:
    payload = context["payload"]
    result = context["result"]
    indicators = context["indicators"]
    drivers = context["drivers"]

    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(0.7)
    section.bottom_margin = Inches(0.7)
    section.left_margin = Inches(0.7)
    section.right_margin = Inches(0.7)

    doc.add_heading("Reporte financiero del programa", level=1)
    doc.add_paragraph(f"Generado: {context['generated_at']}")
    doc.add_paragraph(f"Modelo: {result['model_name']} | modo: {result['mode']}")

    doc.add_heading("Resultado del analisis", level=2)
    table = doc.add_table(rows=4, cols=2)
    table.style = "Table Grid"
    rows = [
        ("Probabilidad de incumplimiento", percent(result["probability"])),
        ("Nivel de riesgo", result["risk_label"].upper()),
        ("Umbral", percent(result["threshold"])),
        ("Decision", "incumplimiento probable" if result["prediction"] else "sin incumplimiento probable"),
    ]
    for index, (label, value) in enumerate(rows):
        table.cell(index, 0).text = label
        table.cell(index, 1).text = value

    doc.add_heading("Indicadores financieros", level=2)
    indicator_table = doc.add_table(rows=1, cols=2)
    indicator_table.style = "Table Grid"
    indicator_table.rows[0].cells[0].text = "Indicador"
    indicator_table.rows[0].cells[1].text = "Valor"
    for label, value in [
        ("Saldos totales", money(indicators["total_bills"])),
        ("Pagos totales", money(indicators["total_payments"])),
        ("Uso promedio de credito", percent(indicators["credit_usage"])),
        ("Cobertura pago / saldo", percent(indicators["payment_coverage"])),
        ("Mora maxima", f"{int(indicators['max_delay'])} meses" if indicators["max_delay"] > 0 else "Sin mora"),
    ]:
        cells = indicator_table.add_row().cells
        cells[0].text = label
        cells[1].text = value

    doc.add_heading("Factores explicativos", level=2)
    if drivers:
        for driver in drivers:
            sign = "+" if driver["direction"] == "increase" else "-"
            doc.add_paragraph(
                f"{driver['factor']}: {sign}{percent(driver['impact'])}. {driver['detail']}",
                style="List Bullet",
            )
    else:
        doc.add_paragraph("El perfil no muestra factores extremos de riesgo.")

    doc.add_heading("Recomendaciones", level=2)
    for item in context["recommendations"]:
        doc.add_paragraph(item, style="List Bullet")

    doc.add_heading("Datos ingresados", level=2)
    data_table = doc.add_table(rows=1, cols=2)
    data_table.style = "Table Grid"
    data_table.rows[0].cells[0].text = "Variable"
    data_table.rows[0].cells[1].text = "Valor"
    for key in FEATURE_COLUMNS:
        cells = data_table.add_row().cells
        cells[0].text = FIELD_LABELS.get(key, key)
        cells[1].text = display_value(key, payload[key])

    buffer = BytesIO()
    doc.save(buffer)
    return buffer.getvalue()


def build_xlsx_report(context: Dict) -> bytes:
    payload = context["payload"]
    result = context["result"]
    indicators = context["indicators"]
    drivers = context["drivers"]
    wb = Workbook()
    ws = wb.active
    ws.title = "Resumen"

    title_fill = PatternFill("solid", fgColor="0A0D14")
    header_fill = PatternFill("solid", fgColor="E8EEF5")
    good_fill = PatternFill("solid", fgColor="DDF7E8")
    warning_fill = PatternFill("solid", fgColor="FFF2CC")
    danger_fill = PatternFill("solid", fgColor="FAD6DB")
    thin_border = Border(bottom=Side(style="thin", color="D9E2EA"))
    bold = Font(bold=True)
    white_bold = Font(bold=True, color="FFFFFF")

    def style_row(sheet, row_number: int, fill: PatternFill = header_fill) -> None:
        for cell in sheet[row_number]:
            cell.font = bold
            cell.fill = fill
            cell.border = thin_border
            cell.alignment = Alignment(vertical="center")

    ws.append(["Reporte financiero del programa", context["generated_at"]])
    style_row(ws, 1, title_fill)
    ws["A1"].font = white_bold
    ws["B1"].font = white_bold
    ws.append(["Modelo", result["model_name"], "Modo", result["mode"]])
    ws.append([])
    ws.append(["Resultado", "Valor"])
    style_row(ws, 4)
    for row in [
        ("Probabilidad de incumplimiento", percent(result["probability"])),
        ("Nivel de riesgo", result["risk_label"].upper()),
        ("Umbral", percent(result["threshold"])),
        ("Decision", "incumplimiento probable" if result["prediction"] else "sin incumplimiento probable"),
    ]:
        ws.append(row)
    risk_row = 6
    ws[f"B{risk_row}"].fill = (
        danger_fill if result["risk_label"] == "alto" else warning_fill if result["risk_label"] == "medio" else good_fill
    )

    ws.append([])
    ws.append(["Indicador", "Valor"])
    style_row(ws, 10)
    for row in [
        ("Saldos totales", money(indicators["total_bills"])),
        ("Pagos totales", money(indicators["total_payments"])),
        ("Uso promedio de credito", percent(indicators["credit_usage"])),
        ("Cobertura pago / saldo", percent(indicators["payment_coverage"])),
        ("Mora maxima", f"{int(indicators['max_delay'])} meses" if indicators["max_delay"] > 0 else "Sin mora"),
    ]:
        ws.append(row)

    ws.append([])
    ws.append(["Factores explicativos", "Impacto"])
    style_row(ws, ws.max_row)
    if drivers:
        for driver in drivers:
            sign = "+" if driver["direction"] == "increase" else "-"
            ws.append([f"{driver['factor']} - {driver['detail']}", f"{sign}{percent(driver['impact'])}"])
            ws.cell(ws.max_row, 2).fill = danger_fill if driver["direction"] == "increase" else good_fill
    else:
        ws.append(["Sin factores extremos detectados", "0%"])

    ws.append([])
    ws.append(["Recomendaciones"])
    style_row(ws, ws.max_row)
    for item in context["recommendations"]:
        ws.append([item])

    data_ws = wb.create_sheet("Datos ingresados")
    data_ws.append(["Variable", "Valor mostrado", "Valor numerico"])
    style_row(data_ws, 1)
    for key in FEATURE_COLUMNS:
        data_ws.append([FIELD_LABELS.get(key, key), display_value(key, payload[key]), payload[key]])
        current_row = data_ws.max_row
        if key.startswith("PAY_AMT"):
            data_ws.cell(current_row, 2).fill = good_fill if float(payload[key]) > 0 else warning_fill
        elif key.startswith("PAY_"):
            data_ws.cell(current_row, 2).fill = danger_fill if float(payload[key]) > 0 else good_fill

    chart_ws = wb.create_sheet("Saldos vs Pagos")
    chart_ws.append(["Mes", "Saldo", "Pago"])
    style_row(chart_ws, 1)
    for month in range(1, 7):
        chart_ws.append([f"M{month}", float(payload[f"BILL_AMT{month}"]), float(payload[f"PAY_AMT{month}"])])
    chart = BarChart()
    chart.title = "Saldos facturados vs pagos"
    chart.y_axis.title = "Soles"
    chart.x_axis.title = "Mes"
    chart.add_data(Reference(chart_ws, min_col=2, max_col=3, min_row=1, max_row=7), titles_from_data=True)
    chart.set_categories(Reference(chart_ws, min_col=1, min_row=2, max_row=7))
    chart_ws.add_chart(chart, "E2")

    for sheet in [ws, data_ws, chart_ws]:
        sheet.freeze_panes = "A2"
        sheet.column_dimensions["A"].width = 38
        sheet.column_dimensions["B"].width = 30
        sheet.column_dimensions["C"].width = 18
        sheet.column_dimensions["D"].width = 18
        for row in sheet.iter_rows():
            for cell in row:
                cell.alignment = Alignment(vertical="top", wrap_text=True)

    buffer = BytesIO()
    wb.save(buffer)
    return buffer.getvalue()


def report_response(content: bytes, filename: str, media_type: str, inline: bool = False) -> Response:
    disposition = "inline" if inline else "attachment"
    return Response(
        content=content,
        media_type=media_type,
        headers={"Content-Disposition": f'{disposition}; filename="{filename}"'},
    )


def normalize_text(value: Any) -> str:
    decomposed = unicodedata.normalize("NFKD", str(value or "").lower())
    return "".join(char for char in decomposed if not unicodedata.combining(char))


def has_any(text: str, phrases: List[str]) -> bool:
    return any(normalize_text(phrase) in text for phrase in phrases)


def retrieve_knowledge(message: str, language: str, limit: int = 4) -> List[Dict[str, str]]:
    text = normalize_text(message)
    scored = []
    for item in ACADEMIC_KNOWLEDGE:
        score = sum(1 for keyword in item["keywords"] if normalize_text(keyword) in text)
        if score:
            scored.append((score, item))
    scored.sort(key=lambda row: row[0], reverse=True)
    return [
        {"id": item["id"], "text": item[language]}
        for _, item in scored[:limit]
    ]


def summarize_prediction_context(prediction_context: Optional[Dict[str, Any]], language: str) -> str:
    if not prediction_context:
        return ""
    probability = prediction_context.get("probability")
    risk_label = prediction_context.get("risk_label")
    model_name = prediction_context.get("model_name")
    drivers = prediction_context.get("drivers") or []
    if probability is None:
        return ""

    try:
        probability_text = percent(float(probability))
    except (TypeError, ValueError):
        probability_text = str(probability)

    driver_text = ""
    if drivers:
        driver_parts = []
        for driver in drivers[:3]:
            label = driver.get("factor") or driver.get("label") or "factor"
            direction = driver.get("direction", "increase")
            direction_text = "aumenta" if direction == "increase" else "reduce"
            if language == "en":
                direction_text = "increases" if direction == "increase" else "reduces"
            impact = driver.get("impact")
            impact_text = f" {percent(float(impact))}" if isinstance(impact, (int, float)) else ""
            driver_parts.append(f"{label} ({direction_text}{impact_text})")
        driver_text = "; ".join(driver_parts)

    if language == "en":
        base = f"Current evaluation: {probability_text} risk, label {risk_label}, model {model_name}."
        return f"{base} Main factors: {driver_text}." if driver_text else base
    base = f"Evaluacion actual: {probability_text} de riesgo, nivel {risk_label}, modelo {model_name}."
    return f"{base} Factores principales: {driver_text}." if driver_text else base


def display_model_name(model_name: Optional[str]) -> str:
    return str(model_name or "LSTM").replace("_", "-")


def normalize_model_name(model_name: Optional[str]) -> str:
    normalized = str(model_name or "LSTM").strip().upper().replace("-", "_").replace(" ", "_")
    aliases = {"CNNLSTM": "CNN_LSTM", "ATTENTION": "LSTM_ATTENTION", "LSTM_ATT": "LSTM_ATTENTION"}
    return aliases.get(normalized, normalized)


def current_context_model(prediction_context: Optional[Dict[str, Any]]) -> str:
    if not prediction_context:
        return "LSTM"
    for key in ["selected_model", "requested_model", "model_name"]:
        if prediction_context.get(key):
            return normalize_model_name(prediction_context.get(key))
    return "LSTM"


def numeric(value: Any) -> Optional[float]:
    try:
        number = float(value)
    except (TypeError, ValueError):
        return None
    return number if number == number else None


def metric(value: Any, digits: int = 3) -> str:
    number = numeric(value)
    return "--" if number is None else f"{number:.{digits}f}"


def comparison_rows(stats: Dict[str, Any]) -> List[Dict[str, Any]]:
    return [row for row in stats.get("comparison", []) if row.get("model")]


def best_row_by_metric(stats: Dict[str, Any], metric_name: str, reverse: bool = True) -> Dict[str, Any]:
    rows = [row for row in comparison_rows(stats) if numeric(row.get(metric_name)) is not None]
    if not rows:
        return {}
    return sorted(rows, key=lambda row: numeric(row.get(metric_name)) or 0.0, reverse=reverse)[0]


def row_for_model(stats: Dict[str, Any], model_name: str) -> Dict[str, Any]:
    normalized = normalize_model_name(model_name)
    for row in comparison_rows(stats):
        if normalize_model_name(row.get("model")) == normalized:
            return row
    return {}


def availability_for_model(model_name: str) -> Dict[str, Any]:
    normalized = normalize_model_name(model_name)
    for model in predictor.available_models():
        if normalize_model_name(model.get("name")) == normalized:
            return model
    return {"name": normalized, "artifact_exists": False, "status": "proxy"}


def context_indicator_text(prediction_context: Optional[Dict[str, Any]], language: str) -> str:
    if not prediction_context:
        return ""
    indicators = prediction_context.get("indicators") or {}
    usage = numeric(indicators.get("usage") or indicators.get("credit_usage"))
    coverage = numeric(indicators.get("coverage") or indicators.get("payment_coverage"))
    max_delay = numeric(indicators.get("maxDelay") or indicators.get("max_delay"))
    parts = []
    if usage is not None:
        parts.append(("usage", percent(usage)))
    if coverage is not None:
        parts.append(("coverage", percent(coverage)))
    if max_delay is not None:
        parts.append(("max delay", f"{max_delay:.0f}"))
    if not parts:
        return ""
    if language == "en":
        labels = {"usage": "credit usage", "coverage": "payment coverage", "max delay": "maximum delay"}
        return ", ".join(f"{labels[key]} {value}" for key, value in parts)
    labels = {"usage": "uso de credito", "coverage": "cobertura de pago", "max delay": "mora maxima"}
    return ", ".join(f"{labels[key]} {value}" for key, value in parts)


def model_recommendation_answer(language: str, stats: Dict[str, Any], prediction_context: Optional[Dict[str, Any]]) -> str:
    production = normalize_model_name(stats.get("production_model") or "LSTM")
    selected = current_context_model(prediction_context)
    auc_row = best_row_by_metric(stats, "roc_auc_mean")
    f1_row = best_row_by_metric(stats, "f1_mean")
    speed_row = best_row_by_metric(stats, "fit_seconds_mean", reverse=False)
    selected_row = row_for_model(stats, selected)
    selected_availability = availability_for_model(selected)

    if language == "en":
        artifact_note = (
            "It has a saved artifact and can be consumed directly by FastAPI."
            if selected_availability.get("artifact_exists")
            else "It is currently available as an academic proxy until its trained artifact is saved."
        )
        return (
            f"For the normal app workflow I recommend {display_model_name(production)}, because it is the production model. "
            f"In the current comparison, the best AUC is {display_model_name(auc_row.get('model'))} "
            f"with AUC {metric(auc_row.get('roc_auc_mean'))}; the strongest F1 candidate is "
            f"{display_model_name(f1_row.get('model'))} with F1 {metric(f1_row.get('f1_mean'))}; and the fastest model is "
            f"{display_model_name(speed_row.get('model'))} with {metric(speed_row.get('fit_seconds_mean'), 2)} seconds. "
            f"Your selected model is {display_model_name(selected)}: AUC {metric(selected_row.get('roc_auc_mean'))}, "
            f"F1 {metric(selected_row.get('f1_mean'))}. {artifact_note} "
            "For a teacher presentation, show all five models, but use the production model for the final prediction and reports."
        )
    artifact_note = (
        "Tiene artefacto guardado y FastAPI lo puede consumir directamente."
        if selected_availability.get("artifact_exists")
        else "Por ahora funciona como proxy academico hasta guardar su artefacto entrenado."
    )
    return (
        f"Para el flujo normal de la app recomiendo {display_model_name(production)}, porque es el modelo en produccion. "
        f"En la comparacion actual, el mejor AUC es {display_model_name(auc_row.get('model'))} "
        f"con AUC {metric(auc_row.get('roc_auc_mean'))}; el mejor candidato por F1 es "
        f"{display_model_name(f1_row.get('model'))} con F1 {metric(f1_row.get('f1_mean'))}; y el mas rapido es "
        f"{display_model_name(speed_row.get('model'))} con {metric(speed_row.get('fit_seconds_mean'), 2)} segundos. "
        f"Tu modelo seleccionado es {display_model_name(selected)}: AUC {metric(selected_row.get('roc_auc_mean'))}, "
        f"F1 {metric(selected_row.get('f1_mean'))}. {artifact_note} "
        "Para exponer ante tu docente, muestra los cinco modelos, pero usa el modelo en produccion para la prediccion final y los reportes."
    )


def purpose_answer(language: str) -> str:
    if language == "en":
        return (
            "This web app is an AI financial advisor for academic and demo use. It estimates the probability of credit default "
            "from a client's profile, credit limit, six months of balances, payments and delay history. It also compares neural "
            "models, explains the main risk drivers, lets users simulate scenarios and generates PDF, Word and Excel reports."
        )
    return (
        "Esta app web es un asesor financiero con IA para uso academico y demostrativo. Estima la probabilidad de incumplimiento "
        "crediticio usando perfil del cliente, limite de credito, seis meses de saldos, pagos e historial de mora. Tambien compara "
        "modelos neuronales, explica los factores principales del riesgo, permite simular escenarios y genera reportes PDF, Word y Excel."
    )


def workflow_answer(language: str) -> str:
    if language == "en":
        return (
            "Use it in this order: 1. enter the client profile, 2. add delay history, balances and payments, 3. choose the model, "
            "4. calculate risk, 5. review the gauge and explanatory factors, 6. compare scenarios if needed, and 7. preview or download reports."
        )
    return (
        "Usala en este orden: 1. ingresa el perfil del cliente, 2. completa mora, saldos y pagos, 3. elige el modelo, "
        "4. calcula el riesgo, 5. revisa el medidor y los factores explicativos, 6. compara escenarios si lo necesitas, "
        "y 7. previsualiza o descarga los reportes."
    )


def risk_answer(language: str, prediction_context: Optional[Dict[str, Any]]) -> str:
    context_summary = summarize_prediction_context(prediction_context, language)
    if not context_summary:
        if language == "en":
            return (
                "The risk score uses payment delays, payment coverage, credit usage, approved limit and profile variables. "
                "Calculate a profile first and then ask me 'why did I get this risk?' for a personalized explanation."
            )
        return (
            "El riesgo usa mora, cobertura de pagos, uso del credito, limite aprobado y variables del perfil. "
            "Primero calcula un perfil y luego preguntame 'por que me dio este riesgo?' para darte una explicacion personalizada."
        )

    probability = numeric(prediction_context.get("probability")) or 0.0
    indicators = context_indicator_text(prediction_context, language)
    drivers = prediction_context.get("drivers") or []
    driver_text = "; ".join(
        (driver.get("detail") or driver.get("factor") or "").rstrip(". ")
        for driver in drivers[:3]
        if driver
    )
    if language == "en":
        action = (
            "This is a high-risk profile: review delays, reduce credit usage and require stronger payment evidence."
            if probability >= 0.70
            else "This is a medium-risk profile: monitor payment coverage and test an optimized scenario."
            if probability >= 0.40
            else "This is a low-risk profile: the profile looks healthier, but still review the explanatory factors."
        )
        return f"{context_summary} Indicators: {indicators or 'not available'}. {action} Main explanation: {driver_text or 'no strong factor detected'}."
    action = (
        "Es un perfil de riesgo alto: revisa mora, reduce uso del credito y exige mayor evidencia de pago."
        if probability >= 0.70
        else "Es un perfil de riesgo medio: monitorea cobertura de pagos y prueba un escenario optimizado."
        if probability >= 0.40
        else "Es un perfil de riesgo bajo: el perfil luce mas saludable, pero igual conviene revisar los factores explicativos."
    )
    return f"{context_summary} Indicadores: {indicators or 'no disponibles'}. {action} Explicacion principal: {driver_text or 'no se detecto un factor fuerte'}."


def statistics_answer(language: str, stats: Dict[str, Any]) -> str:
    tests = stats.get("statistical_tests", [])
    folds = stats.get("required_structure", {}).get("cross_validation_folds", 0)
    auc_row = best_row_by_metric(stats, "roc_auc_mean")
    if language == "en":
        return (
            f"The statistical module validates model comparison with {folds or 5} cross-validation folds, paired t-test and Wilcoxon. "
            f"It currently reads {len(tests)} paired comparisons. The leading AUC model is "
            f"{display_model_name(auc_row.get('model'))} with AUC {metric(auc_row.get('roc_auc_mean'))}. "
            "Use this section to justify whether differences between models are robust or only descriptive."
        )
    return (
        f"El modulo estadistico valida la comparacion con {folds or 5} folds de validacion cruzada, t-test pareado y Wilcoxon. "
        f"Actualmente lee {len(tests)} comparaciones pareadas. El modelo lider por AUC es "
        f"{display_model_name(auc_row.get('model'))} con AUC {metric(auc_row.get('roc_auc_mean'))}. "
        "Usa esta seccion para justificar si las diferencias entre modelos son robustas o solo descriptivas."
    )


def reports_answer(language: str) -> str:
    if language == "en":
        return (
            "The reports belong to the program output, not only to the deliverables. PDF is useful for preview and presentation, "
            "Word for editable academic documentation, and Excel for reviewing the entered data, metrics and charts."
        )
    return (
        "Los reportes son salidas del programa, no solo entregables. El PDF sirve para previsualizar y presentar, "
        "Word para documentacion academica editable, y Excel para revisar datos ingresados, metricas y graficos."
    )


def example_questions_answer(language: str) -> str:
    if language == "en":
        return (
            "Useful questions are: What does this app do? Which model should I use? Why did this profile get this risk? "
            "What do AUC and F1 mean? How do the statistical tests validate the models? What report should I download?"
        )
    return (
        "Preguntas utiles que puedes hacer: Para que sirve esta app? Que modelo debo usar? Por que este perfil dio ese riesgo? "
        "Que significan AUC y F1? Como validan los modelos las pruebas estadisticas? Que reporte debo descargar?"
    )


def chatbot_answer(message: str, language: str, prediction_context: Optional[Dict[str, Any]] = None) -> Dict:
    text = normalize_text(message)
    stats = statistical_validation_summary()
    topics: List[str] = []
    retrieved = retrieve_knowledge(message, language)
    context_used = bool(prediction_context and prediction_context.get("probability") is not None)

    asks_purpose = has_any(text, ["para que", "que hace", "sirve", "programa", "app", "aplicacion", "proyecto"])
    asks_workflow = has_any(text, ["como uso", "como funciona", "pasos", "llenar", "usuario", "usar la app"])
    asks_model = has_any(text, ["modelo", "model", "lstm", "gru", "mlp", "cnn", "attention", "atencion", "debo usar", "recomienda"])
    asks_statistics = has_any(text, ["estadistica", "estadistico", "wilcoxon", "t-test", "ttest", "cross", "fold", "validacion", "prueba"])
    asks_reports = has_any(text, ["reporte", "report", "pdf", "word", "excel", "descargar", "previsualizar"])
    asks_risk = has_any(text, ["porque", "por que", "why", "riesgo", "risk", "resultado", "mora", "pago", "saldo", "incumplimiento"])
    asks_examples = has_any(text, ["preguntas", "usuarios", "que puedo preguntar", "ejemplos", "faq"])

    if asks_examples:
        topics.append("examples")
        answer = example_questions_answer(language)
    elif asks_statistics:
        topics.append("statistics")
        answer = statistics_answer(language, stats)
    elif asks_model:
        topics.append("model_recommendation")
        answer = model_recommendation_answer(language, stats, prediction_context)
    elif asks_risk:
        topics.append("risk")
        answer = risk_answer(language, prediction_context)
    elif asks_reports:
        topics.append("reports")
        answer = reports_answer(language)
    elif asks_workflow:
        topics.append("workflow")
        answer = workflow_answer(language)
    elif asks_purpose:
        topics.append("purpose")
        answer = purpose_answer(language)
    else:
        topics.append("help")
        answer = (
            "I can answer about the app purpose, model choice, risk explanation, reports, statistical validation and deployment. "
            "Ask for example: 'Which model should I use?' or 'Why did this profile get this risk?'"
            if language == "en"
            else "Puedo responder sobre el proposito de la app, eleccion de modelos, explicacion del riesgo, reportes, validacion estadistica y despliegue. "
            "Por ejemplo: 'Que modelo debo usar?' o 'Por que este perfil dio este riesgo?'"
        )

    if retrieved:
        sources_text = " ".join(item["text"] for item in retrieved)
        if language == "en":
            answer = f"{answer}\n\nProject context: {sources_text}"
        else:
            answer = f"{answer}\n\nContexto del proyecto: {sources_text}"

    return {
        "answer": answer,
        "language": language,
        "topics": topics,
        "sources": [item["id"] for item in retrieved],
        "context_used": context_used,
    }


@app.get("/")
def root() -> Dict[str, str]:
    return {
        "name": "Asesor Financiero Personal IA API",
        "status": "ok",
        "docs": "/docs",
    }


@app.get("/health")
def health() -> Dict[str, str]:
    return {"status": "ok", "model_mode": predictor.mode}


@app.get("/model-info")
def model_info() -> Dict:
    return {
        "mode": predictor.mode,
        "features": FEATURE_COLUMNS,
        "metadata": predictor.metadata,
        "available_models": predictor.available_models(),
        "sample_payload": DEFAULT_SAMPLE,
    }


@app.get("/statistics/validation")
def statistics_validation() -> Dict:
    summary = statistical_validation_summary()
    summary["available_models"] = predictor.available_models()
    return summary


@app.post("/chatbot", response_model=ChatbotResponse)
def chatbot(request: ChatbotRequest) -> Dict:
    return chatbot_answer(request.message, request.language, request.prediction_context)


@app.post("/predict", response_model=PredictionResponse)
def predict(request: PredictionRequest) -> Dict:
    return predictor.predict(request.model_dump(), request.model_name)


@app.post("/reports/financial/pdf")
def financial_report_pdf(request: PredictionRequest) -> Response:
    context = report_context(request)
    return report_response(
        build_pdf_report(context),
        "reporte_financiero_programa.pdf",
        "application/pdf",
        inline=True,
    )


@app.post("/reports/financial/docx")
def financial_report_docx(request: PredictionRequest) -> Response:
    context = report_context(request)
    return report_response(
        build_docx_report(context),
        "reporte_financiero_programa.docx",
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    )


@app.post("/reports/financial/xlsx")
def financial_report_xlsx(request: PredictionRequest) -> Response:
    context = report_context(request)
    return report_response(
        build_xlsx_report(context),
        "reporte_financiero_programa.xlsx",
        "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
    )
